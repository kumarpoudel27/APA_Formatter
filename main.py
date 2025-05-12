from fastapi import FastAPI, File, UploadFile, Form
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse, StreamingResponse
from typing import Optional
import uvicorn
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
import io
import requests
import os

app = FastAPI()

# Add CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=[
        "http://localhost:3000",  # Local development
        "https://apa-formatter.vercel.app",  # Vercel frontend
        "https://*.vercel.app"  # Any Vercel subdomain
    ],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# APA formatting functions
APA_FONT = 'Times New Roman'
APA_SIZE = 12
APA_PARAGRAPH_INDENT = 0.5  # inches
APA_LINE_SPACING = 2.0

# Helper for title case (APA headings)
def title_case(s):
    # Capitalize first letter of each major word, APA style (except small words)
    small_words = {'and', 'or', 'the', 'of', 'in', 'on', 'for', 'to', 'a', 'an', 'by', 'at', 'with', 'from'}
    words = s.lower().split()
    return ' '.join([w.capitalize() if i == 0 or w not in small_words else w for i, w in enumerate(words)])

# Helper for sentence case (APA body text)
def sentence_case(s):
    s = s.strip()
    if not s:
        return s
    s = s[0].upper() + s[1:]
    return s

# Helper for APA reference capitalization
def apa_reference_case(s):
    # Capitalize only the first word and proper nouns (simple heuristic)
    s = s.strip()
    if not s:
        return s
    s = s[0].upper() + s[1:]
    return s

# Detect APA headings (improved, stricter)
def detect_heading(text):
    t = re.sub(r'<.*?>', '', text).strip()
    headings = ["abstract", "introduction", "references", "discussion", "conclusion"]
    if t.lower() in headings:
        return 1
    if t.istitle() and len(t.split()) < 10 and t.lower() not in ["keywords:"]:
        return 2
    return 0

# APA 7th reference auto-detection and formatting
def format_apa_reference(ref):
    ref = ref.strip()
    # Try to detect author(s), year, title, source
    # Very basic pattern: Author(s). (Year). Title. Source. DOI/URL
    author_pat = r"^([^.]+?)\\. ?\\((\\d{4})\\)\\. ?"
    match = re.match(author_pat, ref)
    if match:
        authors = match.group(1)
        year = match.group(2)
        rest = ref[match.end():].strip()
        # Title ends at next period
        title_pat = r"([^.]+?)\\. ?"
        title_match = re.match(title_pat, rest)
        if title_match:
            title = title_match.group(1)
            after_title = rest[title_match.end():].strip()
            # Capitalize only first word and proper nouns in title
            def apa_title_case(s):
                s = s.strip()
                if not s:
                    return s
                words = s.split()
                result = [words[0].capitalize()] + [w if w.isupper() else w.lower() for w in words[1:]]
                return ' '.join(result)
            title_fmt = apa_title_case(title)
            # Try to detect journal/book (italicize)
            journal_pat = r"([^.,]+)(,|\\.)"
            journal_match = re.match(journal_pat, after_title)
            if journal_match:
                journal = journal_match.group(1).strip()
                after_journal = after_title[journal_match.end():].strip()
                # Compose APA reference
                apa = f"{authors}. ({year}). <i>{title_fmt}</i>. <i>{journal}</i>"
                if after_journal:
                    apa += f", {after_journal}"
                return apa
            else:
                # No journal detected
                return f"{authors}. ({year}). <i>{title_fmt}</i>. {after_title}"
        else:
            # No title detected
            return ref
    # If not matched, just capitalize first word
    return ref[:1].upper() + ref[1:]

# Hugging Face integration for text cleanup
HF_API_KEY = os.getenv("HF_API_KEY", "")  # Get from environment variable
def huggingface_format_text(text):
    if not HF_API_KEY:
        return text  # Return original text if no API key is set
    API_URL = "https://api-inference.huggingface.co/models/facebook/bart-large-mnli"
    headers = {"Authorization": f"Bearer {HF_API_KEY}"}
    payload = {"inputs": text}
    try:
        response = requests.post(API_URL, headers=headers, json=payload, timeout=20)
        if response.status_code == 200:
            result = response.json()
            # Some models return a list with 'generated_text', others may differ
            if isinstance(result, list) and 'generated_text' in result[0]:
                return result[0]['generated_text']
            elif isinstance(result, dict) and 'generated_text' in result:
                return result['generated_text']
        # Fallback if output format is unexpected
        return text
    except Exception as e:
        return text

def format_apa_reference_with_formatting(ref):
    # Basic APA formatting: sentence case for article, italics for journal/book, italics for volume
    # Returns (formatted_text, formatting_dict)
    # This is a simple heuristic; for perfect results, use a citation parser.
    ref = ref.strip()
    fmt = {'italics': []}
    # Italicize journal/book titles (between "." and ",", or after 'In ')
    # Italicize volume numbers: e.g. 54(3), 15(1), etc.
    # Example: The Journal of Finance, 54(3), 875-899.
    # Find journal/book (after period, before comma, with capital)
    journal_match = re.search(r'\. ([A-Z][^.,]+),', ref)
    if journal_match:
        start, end = journal_match.start(1), journal_match.end(1)
        fmt['italics'].append((start, end))
    # Italicize volume number
    vol_match = re.search(r'(\d+)\(', ref)
    if vol_match:
        start, end = vol_match.start(1), vol_match.end(1)
        fmt['italics'].append((start, end))
    # Sentence case for article title (after year)
    # Not perfect: just lowercase after year and period
    title_match = re.search(r'\(\d{4}\)\. ([^\.]+)\.', ref)
    if title_match is not None and title_match.lastindex:
        # Use group(1) if only one group, group(2) if two groups
        group_num = 2 if title_match.lastindex >= 2 else 1
        title = title_match.group(group_num)
        sentence_case_title = title[:1].upper() + title[1:].lower()
        ref = ref.replace(title, sentence_case_title)
    return ref, fmt

def format_individual_author_name(name_str):
    name_str = name_str.strip()
    if not name_str:
        return ""

    # Heuristic for group authors:
    # e.g., "American Psychological Association", "U.S. Department of Health and Human Services"
    words_in_original = name_str.split()
    is_likely_group = False
    if ',' not in name_str and len(words_in_original) > 1:
        common_org_indicators = [
            "university", "college", "association", "society", "department",
            "center", "institute", "foundation", "corporation", "inc",
            "ltd", "llc", "government", "council", "bureau", "office", "press"
        ]
        # Count words that are fully capitalized (like APA) or title-cased
        cap_words = sum(1 for w in words_in_original if w.istitle() or w.isupper())
        
        if any(ind.lower() in name_str.lower() for ind in common_org_indicators):
            is_likely_group = True
        elif len(words_in_original) > 2 and cap_words >= len(words_in_original) -1 : # Most words are capitalized/title
             is_likely_group = True
        elif len(words_in_original) > 3 and name_str.istitle(): # e.g. The New York Times
            is_likely_group = True


    if is_likely_group:
        return name_str.rstrip('.') + '.' # Ensure it ends with one period

    # If it's already reasonably "Lastname, F. M." or "Lastname, F." - normalize spacing
    match_already_formatted = re.match(r"^([^,]+),\s+((?:[A-Z]\.(?:-[A-Z]\.)?\s*)+)$", name_str)
    if match_already_formatted:
        last_part = match_already_formatted.group(1).strip()
        initial_part = match_already_formatted.group(2).strip()
        initial_part_cleaned = re.sub(r'\s+', ' ', initial_part) # Normalize multiple spaces
        return f"{last_part}, {initial_part_cleaned.rstrip()}"

    last_name = ""
    first_middle_text = ""

    if ',' in name_str: # "Last, First Middle"
        parts = [p.strip() for p in name_str.split(',', 1)]
        last_name = parts[0]
        if len(parts) > 1:
            first_middle_text = parts[1]
    else: # "First Middle Last" or "Last"
        name_tokens = name_str.split()
        if not name_tokens: return ""
        if len(name_tokens) > 1:
            last_name = name_tokens[-1]
            first_middle_text = " ".join(name_tokens[:-1])
        else:
            last_name = name_tokens[0]
            first_middle_text = ""

    if not first_middle_text: # Only a last name
        return last_name.rstrip('.')

    initials = []
    # Process word by word for hyphen handling and multi-char initials like "JM"
    # Ensure space after dots for splitting, then strip periods for processing
    words_in_fm = first_middle_text.replace('.', '. ').split()

    for word_token in words_in_fm:
        word_cleaned = word_token.strip().rstrip('.') # Clean leading/trailing dots for processing this word
        if not word_cleaned: continue

        if '-' in word_cleaned: # Hyphenated names e.g. Jean-Luc -> J.-L.
            hyphenated_parts = word_cleaned.split('-')
            initialed_hyphen_parts = []
            for hp_part in hyphenated_parts:
                hp_part_cleaned = hp_part.strip().rstrip('.')
                if hp_part_cleaned and hp_part_cleaned[0].isalpha():
                    initialed_hyphen_parts.append(hp_part_cleaned[0].upper() + ".")
            if initialed_hyphen_parts:
                initials.append("-".join(initialed_hyphen_parts))
        elif word_cleaned.isupper() and len(word_cleaned) > 1 and word_cleaned.isalpha(): # Handle "JM" -> "J. M."
            for char_initial in word_cleaned:
                initials.append(char_initial.upper() + ".")
        elif word_cleaned and word_cleaned[0].isalpha(): # Standard "John" -> "J."
            initials.append(word_cleaned[0].upper() + ".")
            
    formatted_initials = " ".join(filter(None, initials))

    if not formatted_initials:
        return last_name.rstrip('.')

    return f"{last_name.rstrip('.')}, {formatted_initials}"

def format_apa_reference_with_formatting_and_url(ref):
    # Returns (formatted_text, formatting_dict, url)
    ref = ref.strip()
    fmt = {'italics': []}
    url = None
    # Remove highlight, bold, extra spaces
    ref = re.sub(r'[\u200b\u200c]', '', ref)
    ref = ref.replace('\u202c', '').replace('\u202b', '').strip()

    # --- Author Processing ---
    year_match_for_authors_rule = re.search(r'\s*\((?:\d{4}(?:,\s*[A-Za-z]+(?:\s+\d{1,2})?)?|n\.d\.)\)\.', ref)
    authors_part_original = ""
    date_and_after_original = ref

    if year_match_for_authors_rule:
        authors_part_original = ref[:year_match_for_authors_rule.start()].strip()
        date_and_after_original = ref[year_match_for_authors_rule.start():]
    
    if not authors_part_original:
        pass # No author part to process, ref remains date_and_after_original (or full original ref)
    else:
        temp_author_str = authors_part_original
        temp_author_str = re.sub(r',?\s+et al\.?', '; FakeEtAlAuthorPlaceholder', temp_author_str, flags=re.IGNORECASE)

        raw_author_strings_intermediate = []
        # Try splitting by semicolon first
        split_by_semicolon = [s.strip() for s in temp_author_str.split(';') if s.strip()]

        if len(split_by_semicolon) > 1 or (len(split_by_semicolon) == 1 and split_by_semicolon[0] == "FakeEtAlAuthorPlaceholder"):
             raw_author_strings_intermediate = split_by_semicolon
        else: # No effective semicolon split, or only one actual author string before potential '; FakeEtAlAuthor'
            current_str_to_split = split_by_semicolon[0] if split_by_semicolon else temp_author_str # Use original if split was empty
            
            # Handle "A, B, and C" by temporarily replacing ", and " before splitting by comma
            # Use a very unique placeholder to avoid collision
            placeholder_and = "%%COMMA_AND%%"
            current_str_to_split = current_str_to_split.replace(", and ", f",{placeholder_and}") # Comma before placeholder
            
            # Now split by comma, then restore " and " for the placeholder if it's the last element separator
            split_by_comma = [s.strip() for s in current_str_to_split.split(',') if s.strip()]
            
            # If placeholder_and was used, the item before it is the one before "and"
            # And the item with placeholder_and is the last author.
            # This logic is complex. Simpler: split by " and " first if it's likely the main separator.
            
            # Revised simpler splitting:
            # 1. If " and " appears NOT preceded by a comma, it might be "Author A and Author B".
            # This is hard to distinguish from "A, B and C Company".
            # Let's prioritize splitting by ", " then deal with a final " and Name" if it exists.
            
            author_candidates = []
            # Split by ", and " first to isolate the last author if this pattern exists
            parts_by_comma_and = temp_author_str.split(", and ")
            if len(parts_by_comma_and) > 1 : # e.g. ["A, B", "C"] from "A, B, and C"
                author_candidates.extend([p.strip() for p in parts_by_comma_and[0].split(',') if p.strip()])
                author_candidates.append(parts_by_comma_and[1].strip())
            else: # No ", and ", try splitting by " and " (for "A and B")
                parts_by_and = temp_author_str.split(" and ")
                if len(parts_by_and) > 1: # e.g. ["A", "B"] or ["A, B", "C Company"]
                     author_candidates.extend([p.strip() for p in parts_by_and[0].split(',') if p.strip()])
                     author_candidates.append(parts_by_and[1].strip()) # The part after "and"
                else: # No " and " as primary separator, just split by comma
                    author_candidates.extend([p.strip() for p in temp_author_str.split(',') if p.strip()])
            raw_author_strings_intermediate = author_candidates


        formatted_authors = [format_individual_author_name(name_str) for name_str in raw_author_strings_intermediate]
        formatted_authors = [name for name in formatted_authors if name and name != "FakeEtAlAuthorPlaceholder"]

        num_authors = len(formatted_authors)
        authors_final_str = ""
        if num_authors == 0:
            authors_final_str = authors_part_original # Fallback if parsing failed
        elif num_authors == 1:
            authors_final_str = formatted_authors[0]
        elif num_authors >= 21:
            first_19 = formatted_authors[:19]
            last_author = formatted_authors[-1]
            # Ensure no trailing comma on the 19th author before ellipsis
            authors_final_str = ", ".join(first_19).rstrip(',') + ", ..., " + last_author
        elif num_authors == 2:
            authors_final_str = f"{formatted_authors[0]}, & {formatted_authors[1]}"
        else: # 3 to 20 authors
            authors_final_str = ", ".join(formatted_authors[:-1]).rstrip(',') + ", & " + formatted_authors[-1]
        
        ref = authors_final_str + date_and_after_original
        
    # Find DOI or URL (should not be italicized) - This logic can stay as is
    doi_match = re.search(r'(https?://\S+)', ref)
    if doi_match:
        if doi_match.lastindex:
            url = doi_match.group(1)
            ref = ref.replace(url, '').strip()
        else:
            url = None
    # Replace hyphen in page ranges with en-dash (e.g., 15-26 -> 15–26)
    # Applied after URL removal to avoid affecting URLs.
    # Handles cases like: , 15-26. or (pp. 15-26) or , 15-26)
    ref = re.sub(r'(?<=[,\s(])(\d+)-(\d+)(?=[.\s)]|$)', r'\1–\2', ref)

    # Replace ", and " with ", & " before the last author and year (heuristic for two authors)
    # --- APA Title and Source Formatting ---
    # This section handles sentence casing for titles and italicization for
    # book/report titles, journal titles, and journal volume numbers.

    year_match = re.search(r'\((?:\d{4}(?:,\s*[A-Za-z]+(?:\s+\d{1,2})?)?|n\.d\.)\)\.\s*', ref) # Find (Date). followed by space
    if year_match:
        content_start_after_year = year_match.end()
        text_after_year = ref[content_start_after_year:]
        
        title_segment_match = re.match(r'([^\.]+)\.(.*)', text_after_year)
        
        if title_segment_match:
            extracted_title_text = title_segment_match.group(1).strip()
            remaining_text_after_title = title_segment_match.group(2).strip()

            words = extracted_title_text.split()
            if words:
                words[0] = words[0].capitalize()
                processed_words = [words[0]]
                for i, w in enumerate(words[1:], 1):
                    if w.isupper(): # Keep acronyms
                        processed_words.append(w)
                    elif w.istitle(): # Keep already title-cased (proper nouns/subtitles heuristic)
                        processed_words.append(w)
                    else:
                        processed_words.append(w.lower())
                sentence_cased_title = ' '.join(processed_words)
            else:
                sentence_cased_title = extracted_title_text
            
            ref = ref.replace(extracted_title_text, sentence_cased_title, 1)
            
            title_start_in_ref = content_start_after_year
            title_end_in_ref = title_start_in_ref + len(sentence_cased_title)

            journal_pattern = re.compile(
                r'^\s*([A-Z][^,(]+?)'                # Group 1: Journal Title (non-greedy)
                r',?\s*(\d+)'                         # Group 2: Volume (italicized)
                r'(?:\s*\((\w+)\))?'                  # Group 3: Issue (optional, in parens, not italicized)
                r',?\s*([\d–]+(?:–[\d–]+)?)?\.'       # Group 4: Pages (optional, not italicized, uses en-dash)
            )
            journal_match_obj = journal_pattern.match(remaining_text_after_title)
            is_article = bool(journal_match_obj)

            if is_article:
                journal_title_str = journal_match_obj.group(1).strip()
                journal_volume_str = journal_match_obj.group(2).strip()
                try:
                    j_title_start = ref.index(journal_title_str, title_end_in_ref)
                    j_title_end = j_title_start + len(journal_title_str)
                    fmt['italics'].append((j_title_start, j_title_end))

                    j_vol_start = ref.index(journal_volume_str, j_title_end)
                    j_vol_end = j_vol_start + len(journal_volume_str)
                    fmt['italics'].append((j_vol_start, j_vol_end))
                except ValueError:
                    pass
            else: # Book/Report
                title_proper_match = re.match(r"(.+?)(?:\s*(\((?:(?:[Nn]o|Vol)\.\s*\d+|[2-9]|[1-9]\d+)th\s*ed\.|(?:[Rr]evised|[Ss]pecial)\s*ed\.|(?:[Tt]ech(?:nical)?\s*[Rr]ep(?:ort)?(?:\s*[Nn]o.)?)\s*[^)]+\)))$", sentence_cased_title)
                title_to_italicize_text = sentence_cased_title
                if title_proper_match and title_proper_match.group(1):
                    title_to_italicize_text = title_proper_match.group(1).strip()
                
                italic_title_end = title_start_in_ref + len(title_to_italicize_text)
                if title_start_in_ref < italic_title_end:
                     fmt['italics'].append((title_start_in_ref, italic_title_end))
    # Clean up extra spaces and punctuation
    ref = re.sub(r'\s+,', ',', ref)
    ref = re.sub(r'\s+\.', '.', ref)
    ref = re.sub(r'\s+', ' ', ref)
    ref = ref.strip()
    return ref, fmt, url

@app.post("/format-apa/")
async def format_apa(
    text: Optional[str] = Form(None),
    file: Optional[UploadFile] = File(None),
    output_format: Optional[str] = Form("text")
):
    if file:
        doc_in = Document(file.file)
        content = "\n".join([p.text for p in doc_in.paragraphs])
        # Extract tables and images
        extracted_tables = doc_in.tables
        extracted_images = []
        for rel in doc_in.part.rels.values():
            if "image" in rel.target_ref:
                extracted_images.append(rel.target_part.blob)
    elif text:
        content = text
        extracted_tables = []
        extracted_images = []
    else:
        return JSONResponse({"error": "No input provided."}, status_code=400)

    # Clean up the text using Hugging Face before APA formatting
    content = huggingface_format_text(content)

    # Split into lines/paragraphs
    nonempty_lines = []
    for line in content.split('\n'):
        stripped = line.strip()
        if stripped:
            nonempty_lines.append(stripped)
    # Use up to 7 non-empty lines for the title page (APA allows department/school as a line)
    TITLE_PAGE_LINES = 7
    title_block = [nonempty_lines[i] if i < len(nonempty_lines) else "" for i in range(TITLE_PAGE_LINES)]
    body_and_refs = nonempty_lines[TITLE_PAGE_LINES:] if len(nonempty_lines) > TITLE_PAGE_LINES else []
    # Find references section (case-insensitive, first line that starts with 'references')
    ref_idx = -1
    for i, p in enumerate(body_and_refs):
        if p.lower().startswith("references"):
            ref_idx = i
            break
    if ref_idx != -1:
        main_body = body_and_refs[:ref_idx]
        references = body_and_refs[ref_idx:]
    else:
        main_body = body_and_refs
        references = []

    # If output_format is 'docx', return a downloadable .docx file
    if output_format == 'docx':
        doc = Document()
        style = doc.styles['Normal']
        font = style.font
        font.name = APA_FONT
        font.size = Pt(APA_SIZE)
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        # Helper to set font for all runs in a paragraph
        def set_apa_font(para):
            for run in para.runs:
                run.font.name = APA_FONT
                run.font.size = Pt(APA_SIZE)
        # --- TITLE PAGE ---
        section = doc.sections[0]
        section.start_type = WD_SECTION.NEW_PAGE
        # Set APA-compliant margins
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        for _ in range(3):
            doc.add_paragraph("")
        title_page_fields = [
            title_case(title_block[0]) if len(title_block) > 0 else "",
            sentence_case(title_block[1]) if len(title_block) > 1 else "",
            sentence_case(title_block[2]) if len(title_block) > 2 else "",
            sentence_case(title_block[3]) if len(title_block) > 3 else "",
            sentence_case(title_block[4]) if len(title_block) > 4 else "",
            sentence_case(title_block[5]) if len(title_block) > 5 else "",
            sentence_case(title_block[6]) if len(title_block) > 6 else ""
        ]
        for i, field in enumerate(title_page_fields):
            para = doc.add_paragraph(field)
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.line_spacing = APA_LINE_SPACING
            para.paragraph_format.first_line_indent = 0
            if i == 0 and para.runs:
                para.runs[0].bold = True
            set_apa_font(para)
        for section in doc.sections:
            header = section.header
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn
            paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = paragraph.add_run()
            fldChar = OxmlElement('w:fldChar')
            fldChar.set(qn('w:fldCharType'), 'begin')
            instrText = OxmlElement('w:instrText')
            instrText.text = 'PAGE'
            fldChar2 = OxmlElement('w:fldChar')
            fldChar2.set(qn('w:fldCharType'), 'end')
            run._r.append(fldChar)
            run._r.append(instrText)
            run._r.append(fldChar2)
            run.font.size = Pt(APA_SIZE)
            run.font.name = APA_FONT
        # --- ABSTRACT PAGE ---
        # Always add a page break after the title page
        doc.add_page_break()
        abstract_handled = False
        # Detect and render abstract if present, else treat all as main body
        if main_body and main_body[0].strip().lower() == "abstract":
            # Abstract heading: centered, bold, no indent
            para = doc.add_paragraph("Abstract")
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.first_line_indent = 0
            para.paragraph_format.line_spacing = APA_LINE_SPACING
            set_apa_font(para)
            # Ensure all runs in the heading are bold
            for run in para.runs:
                run.bold = True
            # Abstract text: left-aligned, no indent, double-spaced, max 250 words, single paragraph
            if len(main_body) > 1:
                abstract_text = main_body[1].strip()
                # Remove extra newlines, ensure single paragraph
                abstract_text = re.sub(r'\s*\n+\s*', ' ', abstract_text)
                # Limit to 250 words (APA max)
                words = abstract_text.split()
                if len(words) > 250:
                    abstract_text = ' '.join(words[:250]) + '...'
                para = doc.add_paragraph(abstract_text)
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                para.paragraph_format.first_line_indent = 0
                para.paragraph_format.line_spacing = APA_LINE_SPACING
                set_apa_font(para)
            # Keywords (optional, if present)
            if len(main_body) > 2 and main_body[2].lower().startswith("keywords:"):
                para = doc.add_paragraph(main_body[2])
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                para.paragraph_format.first_line_indent = Inches(0.5)
                para.paragraph_format.line_spacing = APA_LINE_SPACING
                set_apa_font(para)
            main_body = main_body[3:] if len(main_body) > 2 and main_body[2].lower().startswith("keywords:") else main_body[2:]
            abstract_handled = True
        # --- MAIN BODY ---
        # If any content remains after title page, always add it after the break
        if main_body:
            if abstract_handled:
                doc.add_page_break()
            # Repeat title at top of first body page, centered, bold, double-spaced
            para = doc.add_paragraph(title_case(title_block[0]))
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.first_line_indent = 0
            para.paragraph_format.line_spacing = APA_LINE_SPACING
            set_apa_font(para)
            if para.runs:
                para.runs[0].bold = True
            # Body paragraphs and headings
            body_idx = 0
            for p in main_body:
                if not p.strip():
                    continue
                if p.strip().lower() == "introduction":
                    continue
                # Insert tables and images at their original positions (approximate)
                # Insert table if this paragraph matches a table marker
                if extracted_tables and body_idx < len(extracted_tables):
                    # If the original doc had a table at this position, insert it
                    # (Assume a placeholder like [TABLE] or just insert after every N paragraphs for demo)
                    table = extracted_tables[body_idx]
                    rows = len(table.rows)
                    cols = len(table.columns)
                    new_table = doc.add_table(rows=rows, cols=cols)
                    new_table.style = 'Table Grid'
                    for i in range(rows):
                        for j in range(cols):
                            new_table.cell(i, j).text = table.cell(i, j).text
                    body_idx += 1
                # Insert image if available (for demo, insert after every N paragraphs)
                if extracted_images and body_idx < len(extracted_images):
                    img_bytes = extracted_images[body_idx]
                    img_path = f"temp_img_{body_idx}.png"
                    with open(img_path, "wb") as f:
                        f.write(img_bytes)
                    doc.add_picture(img_path, width=Inches(4.5))
                    os.remove(img_path)
                    body_idx += 1
                level = detect_heading(p)
                if level == 1:
                    para = doc.add_paragraph(title_case(p))
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    para.paragraph_format.first_line_indent = 0
                    para.paragraph_format.line_spacing = APA_LINE_SPACING
                    set_apa_font(para)
                    if para.runs:
                        para.runs[0].bold = True
                elif level == 2:
                    para = doc.add_paragraph(title_case(p))
                    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    para.paragraph_format.first_line_indent = 0
                    para.paragraph_format.line_spacing = APA_LINE_SPACING
                    set_apa_font(para)
                    if para.runs:
                        para.runs[0].bold = True
                elif level == 3:
                    para = doc.add_paragraph(title_case(p))
                    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    para.paragraph_format.first_line_indent = 0
                    para.paragraph_format.line_spacing = APA_LINE_SPACING
                    set_apa_font(para)
                    if para.runs:
                        para.runs[0].bold = True
                        para.runs[0].italic = True
                elif p.lower().startswith("keywords:"):
                    # Should not appear in body, skip
                    continue
                else:
                    para = doc.add_paragraph(sentence_case(p))
                    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    para.paragraph_format.first_line_indent = Inches(APA_PARAGRAPH_INDENT)
                    para.paragraph_format.line_spacing = APA_LINE_SPACING
                    set_apa_font(para)
        # --- REFERENCES ---
        if references:
            print("DEBUG references:", references)
            doc.add_page_break()
            # Always use 'References' heading, centered, bold, no colon
            para = doc.add_paragraph("References")
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.first_line_indent = 0
            para.paragraph_format.line_spacing = APA_LINE_SPACING
            set_apa_font(para)
            if para.runs:
                para.runs[0].bold = True
            # Remove any line that starts with 'Reference' or 'References' (case-insensitive, with or without colon/whitespace)
            ref_entries = [r for r in references if not re.match(r'^\s*references?:?\s*$', r.strip(), re.IGNORECASE) and r.strip()]
            # Sort entries alphabetically by first author last name (ignore 'A', 'An', 'The')
            # Sort entries alphabetically by first author last name or title
            def apa_sort_key(entry_str):
                entry_str = entry_str.strip()
                
                # Attempt to extract author part before date like (2020). or (n.d.).
                year_match = re.search(r'\s*\((?:\d{4}(?:,\s*[A-Za-z]+(?:\s+\d{1,2})?)?|n\.d\.)\)\.', entry_str)
                sortable_part = entry_str
                if year_match:
                    sortable_part = entry_str[:year_match.start()].strip()

                # Remove leading articles "A ", "An ", "The " for group authors or titles
                leading_articles_pattern = r'^(?:A|An|The)\s+'
                if re.match(leading_articles_pattern, sortable_part, re.IGNORECASE):
                    if ',' not in sortable_part.split('.')[0]: # Heuristic: if no "Last, F."
                         sortable_part = re.sub(leading_articles_pattern, '', sortable_part, count=1, flags=re.IGNORECASE)
                
                author_match = re.match(r'^([^,(]+),?\s*(?:[A-Z]\.|\(etal\))', sortable_part)
                if author_match:
                    return author_match.group(1).lower()

                first_word_match = re.match(r'^([\w\-]+)', sortable_part)
                if first_word_match:
                    return first_word_match.group(1).lower()
                
                return sortable_part.lower() # Fallback

            ref_entries = sorted(ref_entries, key=apa_sort_key)
            
            # Add each reference entry, formatted
            for ref_text_original in ref_entries:
                formatted_ref_text, formatting_info, ref_url = format_apa_reference_with_formatting_and_url(ref_text_original)
                
                para = doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                para.paragraph_format.left_indent = Inches(0.5)
                para.paragraph_format.first_line_indent = Inches(-0.5) # Hanging indent
                para.paragraph_format.line_spacing = APA_LINE_SPACING

                current_pos = 0
                # Ensure 'italics' key exists and is a list
                sorted_italics = sorted(formatting_info.get('italics', []), key=lambda x: x[0])


                for start_italic, end_italic in sorted_italics:
                    if start_italic > current_pos:
                        para.add_run(formatted_ref_text[current_pos:start_italic])
                    # Add italic segment, ensuring indices are valid
                    if start_italic < end_italic <= len(formatted_ref_text):
                        run = para.add_run(formatted_ref_text[start_italic:end_italic])
                        run.italic = True
                    current_pos = end_italic
                
                if current_pos < len(formatted_ref_text):
                    para.add_run(formatted_ref_text[current_pos:])

                if ref_url:
                    space_needed = " "
                    if not formatted_ref_text or formatted_ref_text.endswith(" ") or formatted_ref_text.endswith("."):
                         space_needed = ""
                    run = para.add_run(f"{space_needed}{ref_url}")

                set_apa_font(para)
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return StreamingResponse(buf, media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document', headers={"Content-Disposition": "attachment; filename=formatted_apa.docx"})

    # Otherwise, return formatted text (APA style)
    formatted = []
    # Title page block
    for i, p in enumerate(title_block):
        if i == 0:
            formatted.append(f'<div style="text-align:center;font-weight:bold;font-size:1.1em;margin-bottom:0.2em;line-height:2">{title_case(p)}</div>')
        else:
            formatted.append(f'<div style="text-align:center;font-size:1.1em;margin-bottom:0.2em;line-height:2">{sentence_case(p)}</div>')
    formatted.append('<br/>')
    # Main body
    for p in main_body:
        if p.lower().startswith("keywords:"):
            formatted.append(f"<div style='margin-left:2em;margin-bottom:1em;line-height:2'><i>{sentence_case(p)}</i></div>")
            continue
        level = detect_heading(p)
        if level == 1:
            formatted.append(f"<div style='text-align:center;font-weight:bold;margin-top:1em;margin-bottom:0.3em;line-height:2'>{title_case(p)}</div>")
        elif level == 2:
            formatted.append(f"<div style='font-weight:bold;margin-top:1em;margin-bottom:0.3em;line-height:2'>{title_case(p)}</div>")
        else:
            formatted.append(f"<div style='text-indent:2em;margin-bottom:0.5em;line-height:2'>{sentence_case(p)}</div>")
    # References
    if references:
        for i, p in enumerate(references):
            if i == 0:
                formatted.append(f"<div style='text-align:center;font-weight:bold;margin-top:1em;margin-bottom:0.3em;line-height:2'>{title_case(p)}</div>")
            else:
                apa_ref, fmt, url = format_apa_reference_with_formatting_and_url(p)
                formatted.append(f"<div style='margin-left:2em;text-indent:-2em;margin-bottom:0.3em;line-height:2'>{apa_ref} {url}</div>")
    return {"formatted": "\n".join(formatted)}

if __name__ == "__main__":
    uvicorn.run(app, host="0.0.0.0", port=8000)
